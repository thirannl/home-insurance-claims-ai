from datetime import datetime
import json
import requests
from docx import Document

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings
from langchain_core.documents import Document as LC_Document


# ----------------------------
# READ DOCX
# ----------------------------
def read_docx(path):
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs])


# ----------------------------
# LOAD CLAIM
# ----------------------------
def load_claim(path):
    with open(path, "r") as f:
        return json.load(f)

def validate_claim(claim):
    errors = []

    required_fields = [
        "incident_type",
        "description",
        "estimated_value",
        "date_of_incident"
    ]

    # Check missing / empty
    for field in required_fields:
        if field not in claim or claim[field] in [None, ""]:
            errors.append(f"Missing {field}")

    # Stop early if missing fields
    if errors:
        return errors

    # Type validation
    try:
        value = float(claim["estimated_value"])
        if value <= 0:
            errors.append("estimated_value must be greater than 0")
    except:
        errors.append("estimated_value must be a number")

    # Description check
    if len(claim["description"]) < 10:
        errors.append("description too short")

    # Date format check
    try:
        datetime.strptime(claim["date_of_incident"], "%Y-%m-%d")
    except:
        errors.append("date_of_incident must be YYYY-MM-DD")

    return errors

# ----------------------------
# CREATE VECTOR DB
# ----------------------------
def create_vector_db(text, collection_name):

    # Convert to LangChain document
    docs = [LC_Document(page_content=text)]

    # Split
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=80
    )
    chunks = splitter.split_documents(docs)

    # Embeddings
    embeddings = OllamaEmbeddings(model="nomic-embed-text")

    # Create DB
    db = Chroma.from_documents(
        chunks,
        embedding=embeddings,
        collection_name=collection_name
    )

    return db


# ----------------------------
# RETRIEVE CONTEXT
# ----------------------------
def retrieve_context(db, query, k=4):
    docs = db.similarity_search(query, k=k)
    return "\n\n".join([d.page_content for d in docs])


# ----------------------------
# CLEAN JSON
# ----------------------------
def extract_json(text):
    start = text.find("{")
    end = text.rfind("}") + 1

    if start != -1 and end != -1:
        return text[start:end]
    return None


# ----------------------------
# MAIN
# ----------------------------
def main():

    # Load inputs
    claim = load_claim("data/claim.json")
    errors = validate_claim(claim)

    if errors:
        print(json.dumps({
            "decision": "Needs Human Review",
            "justification": "Invalid or missing claim data",
            "flags": errors
        }, indent=2))
        return
    policy_text = read_docx("data/sample-policy using llm.docx")
    terms_text = read_docx("data/sample-t&c.docx")

    claim_text = json.dumps(claim, indent=2)

    # Build query for retrieval
    query = f"""
    Incident: {claim['incident_type']}
    Description: {claim['description']}
    Amount: {claim['estimated_value']}
    """

    # Create vector DBs
    policy_db = create_vector_db(policy_text, "policy")
    terms_db = create_vector_db(terms_text, "terms")

    # Retrieve relevant content
    policy_context = retrieve_context(policy_db, query)
    terms_context = retrieve_context(terms_db, query)

    print("\n===== POLICY CONTEXT =====\n")
    print(policy_context[:800])

    print("\n===== TERMS CONTEXT =====\n")
    print(terms_context[:800])

    # 🔥 YOUR SAME PROMPT (kept intact)
    prompt = f"""
You are an expert insurance claims assessor.

Your task is to evaluate a claim using ONLY the given policy and terms.

----------------------------
STRICT RULES
----------------------------
1. Use ONLY the given claim, policy, and terms.
2. Do NOT assume anything not written.
3. Do NOT modify the claim.
4. Use ONLY relevant sections of policy.
5. Check numerical calculations carefully.

----------------------------
CRITICAL DECISION RULES
----------------------------
1. Exclusions override everything.
2. Limits override coverage.
3. You MUST check limits before final decision.
4. If claim exceeds limit AND endorsement is not provided:
   → decision MUST be "Needs Human Review"
5. Do NOT return "Covered" if limits are violated.

----------------------------
DECISION FLOW
----------------------------
Step 1: Identify incident type  
Step 2: Check coverage  
- Identify the relevant coverage section for the incident type  
- If the incident type is covered BUT:
    • required condition cannot be verified or is missing → RETURN "Needs Human Review" (STOP)
    • all required conditions in the clause are satisfied → proceed to next step
    -Check any negligence is applied for description of the claim -> Return "Not Covered" if negligence is found
  

Step 3: Check exclusions  
Step 4: Check limits (CRITICAL)
    - If exceeded:
        → Check endorsement
        → If not mentioned:
            → Needs Human Review (STOP)

Step 5: Check conditions  
Step 6: Final decision  

----------------------------
INPUT
----------------------------
Claim:
{claim_text}

Policy:
{policy_context}

Terms:
{terms_context}

----------------------------
OUTPUT (STRICT JSON ONLY)
----------------------------
Then provide ONLY ONE final JSON block at the end.

IMPORTANT:
- JSON must appear ONLY ONCE
- JSON must be the LAST thing in the response
- Do NOT repeat JSON
- Do NOT add anything after JSON


Rules for flags:
- flags should ALWAYS contain meaningful information
- If claim is Covered:
    → include conditions to proceed (e.g., deductible, documents required)
- If claim is Not Covered:
    → include the exact reason (e.g., exclusion triggered)
- If claim is Needs Human Review:
    → include missing or unclear information

- Do NOT leave flags empty
- Do NOT add irrelevant flags

{{
  "decision": "Covered | Not Covered | Needs Human Review",
  "justification": "Short explanation based on policy",
  "flags": ["meaningful conditions or issues"]
}}
"""

    # Call Ollama
    response = requests.post(
        "http://localhost:11434/api/generate",
        json={
            "model": "llama3.1:8b",
            "prompt": prompt,
            "stream": False
        }
    )

    raw_output = response.json()["response"]

    print("\n===== RAW OUTPUT =====\n")
    print(raw_output)

    # Extract JSON
    cleaned = extract_json(raw_output)

    if not cleaned:
        print("\n⚠️ JSON extraction failed")
        return

    try:
        parsed = json.loads(cleaned)

        print("\n===== FINAL OUTPUT =====\n")
        print(json.dumps(parsed, indent=2, ensure_ascii=False))

    except:
        print("\n⚠️ JSON parsing failed")


# ----------------------------
# RUN
# ----------------------------
if __name__ == "__main__":
    main()